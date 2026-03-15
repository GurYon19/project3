from docx import Document

doc_old = Document('Part1-2_Final_Report_Formatted.docx')
doc_new = Document()
doc_new.add_paragraph('Test text before image')
doc_new.element.body.append(doc_old.paragraphs[56]._element)
doc_new.add_paragraph('Test text after image')
doc_new.save('test_move.docx')
print("Test move successful")
