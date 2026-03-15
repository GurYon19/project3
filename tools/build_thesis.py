import re
from docx import Document
from docx.shared import Pt, RGBColor

def build_thesis(md_file, output_docx):
    doc_new = Document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            doc_new.add_paragraph()
            continue
            
        # Handle Custom Image placements
        if line.startswith('[IMAGE_'):
            idx = line.replace('[IMAGE_', '').replace(']', '')
            p = doc_new.add_paragraph(f"[PLACEHOLDER: PASTE IMAGE {idx} HERE]")
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
            continue
            
        # Handle Headings
        if line.startswith('# '):
            doc_new.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc_new.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc_new.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            p = doc_new.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line[2:])
        elif len(line) > 0 and line[0].isdigit() and '. ' in line[:3]:
            p = doc_new.add_paragraph(style='List Number')
            _add_formatted_text(p, line[line.find('. ')+2:])
        else:
            p = doc_new.add_paragraph()
            _add_formatted_text(p, line)
            
    # Apply unifying styles
    for para in doc_new.paragraphs:
        style_name = para.style.name.lower()
        if 'heading 1' in style_name:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(47, 84, 150)
                run.bold = True
        elif 'heading 2' in style_name:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(47, 84, 150)
                run.bold = True
        elif 'heading 3' in style_name:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(31, 55, 99)
                run.bold = True
        else:
            for run in para.runs:
                # Assuming normal formatting
                if not 'w:drawing' in para._element.xml and run.font.color.rgb != RGBColor(255, 0, 0):
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)

    doc_new.save(output_docx)
    print("Thesis document compiled successfully!")

def _add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*') and len(subpart) > 2:
                    run = paragraph.add_run(subpart[1:-1])
                    run.italic = True
                elif subpart.startswith('`') and subpart.endswith('`') and len(subpart) > 2:
                    run = paragraph.add_run(subpart[1:-1])
                    run.font.name = 'Consolas'
                else:
                    paragraph.add_run(subpart)

if __name__ == "__main__":
    build_thesis('thesis_report.md', 'Part1-2_Final_Report_Formatted.docx')
